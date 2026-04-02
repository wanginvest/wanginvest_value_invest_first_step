# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
# ]
# ///

"""Generate a readable DOCX report from MR Dang stock analysis JSON data."""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


# ── Color palette ──────────────────────────────────────────────────────────────
COLOR_TITLE = RGBColor(0x1A, 0x37, 0x6C)   # dark navy
COLOR_SECTION = RGBColor(0x2E, 0x75, 0xB6)  # blue
COLOR_PASS = RGBColor(0x37, 0x86, 0x44)     # green
COLOR_WARN = RGBColor(0xBF, 0x8F, 0x00)     # amber
COLOR_FAIL = RGBColor(0xC0, 0x00, 0x00)     # red
COLOR_GRAY = RGBColor(0x59, 0x59, 0x59)     # gray


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background color."""
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = COLOR_TITLE if level == 1 else COLOR_SECTION
    run.font.bold = True


def _add_kv_table(doc: Document, rows: list[tuple[str, str, str | None]]) -> None:
    """Add a key-value table. rows = [(key, value, status_color_hex|None)]"""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["指标", "数值", "说明"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(hdr[i], "D9E1F2")

    for key, value, color_hex in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)
        if color_hex:
            _set_cell_bg(row[1], color_hex)

    doc.add_paragraph()


def _status_color(status: str) -> str:
    """Map status string to hex color."""
    s = status.lower()
    if any(x in s for x in ["通过", "达标", "稳定", "pass"]):
        return "C6EFCE"
    if any(x in s for x in ["存疑", "提示", "warn", "注意"]):
        return "FFEB9C"
    if any(x in s for x in ["淘汰", "不达标", "危险", "fail", "高风险"]):
        return "FFC7CE"
    return "FFFFFF"


def build_report(data: dict, output_path: Path) -> None:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover / Title ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("MR Dang 价值选股打分报告")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITLE

    meta = data.get("meta", {})
    stock_name = meta.get("stock_name", data.get("stock_name", "未知"))
    ts_code = meta.get("ts_code", data.get("ts_code", ""))
    industry = meta.get("industry", data.get("industry", ""))
    analysis_date = meta.get("analysis_date", data.get("analysis_date", datetime.now().strftime("%Y-%m-%d")))
    data_source = meta.get("data_source", data.get("data_source", ""))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"{stock_name}（{ts_code}）  |  {industry}  |  {analysis_date}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = COLOR_GRAY

    if data_source:
        src_p = doc.add_paragraph()
        src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        src_run = src_p.add_run(f"数据来源：{data_source}")
        src_run.font.size = Pt(10)
        src_run.font.color.rgb = COLOR_WARN

    doc.add_paragraph()

    # ── Section 1: 基础筛查 ───────────────────────────────────────────────────
    _add_heading(doc, "一、基础筛查结果", level=1)
    screening = data.get("screening", [])
    if screening:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["筛查项", "结果", "说明"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
            _set_cell_bg(hdr[i], "D9E1F2")
        for item in screening:
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            status = item.get("result", "")
            row[1].text = status
            _set_cell_bg(row[1], _status_color(status))
            row[2].text = item.get("note", "")
        doc.add_paragraph()

    conclusion = data.get("screening_conclusion", "")
    if conclusion:
        p = doc.add_paragraph()
        p.add_run("筛查结论：").bold = True
        p.add_run(conclusion)
    doc.add_paragraph()

    # ── Section 2: 核心数据 ───────────────────────────────────────────────────
    _add_heading(doc, "二、核心数据概览", level=1)

    valuation = data.get("valuation", {})
    if valuation:
        _add_heading(doc, "估值指标", level=2)
        rows = [
            ("PE(TTM)", valuation.get("pe_ttm", "N/A"), None),
            ("PB", valuation.get("pb", "N/A"), None),
            ("总市值", valuation.get("total_mv", "N/A"), None),
            ("流通市值", valuation.get("circ_mv", "N/A"), None),
        ]
        _add_kv_table(doc, rows)

    financials = data.get("financials", {})
    if financials:
        _add_heading(doc, "财务指标", level=2)
        rows = [
            ("股息率(TTM)", financials.get("dv_ratio", "N/A"), None),
            ("资产负债率", financials.get("debt_to_assets", "N/A"), None),
            ("ROE", financials.get("roe", "N/A"), None),
            ("经营现金流/股", financials.get("ocfps", "N/A"), None),
            ("基本EPS", financials.get("basic_eps", "N/A"), None),
            ("扣非EPS", financials.get("dt_eps", "N/A"), None),
        ]
        _add_kv_table(doc, rows)

    dividend = data.get("dividend", {})
    if dividend:
        _add_heading(doc, "分红历史", level=2)
        rows = [
            ("近3年分红次数", dividend.get("dividend_count", "N/A"), None),
            ("平均每10股派息(元)", dividend.get("avg_cash_div_per_10_shares", "N/A"), None),
            ("分红稳定性", dividend.get("dividend_stability", "N/A"), None),
        ]
        _add_kv_table(doc, rows)

    business = data.get("business", {})
    if business:
        _add_heading(doc, "业务概况", level=2)
        for label, key in [("主营业务", "main_business"), ("行业地位", "industry_position"), ("资源/成本优势", "cost_advantage")]:
            val = business.get(key, "")
            if val:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{label}：").bold = True
                p.add_run(val)
        doc.add_paragraph()

    # ── Section 3: 打分明细 ───────────────────────────────────────────────────
    _add_heading(doc, "三、维度打分明细", level=1)
    scores = data.get("scores", [])
    if scores:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["维度", "得分", "满分", "评分依据"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
            _set_cell_bg(hdr[i], "D9E1F2")
        for item in scores:
            row = table.add_row().cells
            row[0].text = item.get("dimension", "")
            score = item.get("score", 0)
            full = item.get("full_score", 0)
            row[1].text = str(score)
            row[2].text = str(full)
            row[3].text = item.get("reason", "")
            # Color score cell by ratio
            ratio = score / full if full else 0
            if ratio >= 0.7:
                _set_cell_bg(row[1], "C6EFCE")
            elif ratio >= 0.4:
                _set_cell_bg(row[1], "FFEB9C")
            else:
                _set_cell_bg(row[1], "FFC7CE")
        doc.add_paragraph()

    total_score = data.get("total_score", "")
    rating = data.get("rating", "")
    if total_score or rating:
        p = doc.add_paragraph()
        p.add_run(f"总分：{total_score} / 100　　评级：{rating}").bold = True
    doc.add_paragraph()

    # ── Section 4: 操作建议 ───────────────────────────────────────────────────
    _add_heading(doc, "四、操作建议", level=1)
    suggestion = data.get("suggestion", "")
    if suggestion:
        doc.add_paragraph(suggestion)
    doc.add_paragraph()

    # ── Section 5: 买入前清单 ─────────────────────────────────────────────────
    _add_heading(doc, "五、买入前清单核验", level=1)
    checklist = data.get("checklist", [])
    if checklist:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["清单项", "状态", "说明"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
            _set_cell_bg(hdr[i], "D9E1F2")
        for item in checklist:
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            status = item.get("status", "")
            row[1].text = status
            _set_cell_bg(row[1], _status_color(status))
            row[2].text = item.get("note", "")
        doc.add_paragraph()

    passed = sum(1 for i in checklist if "达标" in i.get("status", ""))
    if checklist:
        p = doc.add_paragraph()
        p.add_run(f"达标项：{passed} / {len(checklist)}").bold = True
    doc.add_paragraph()

    # ── Section 6: 财务风险 ───────────────────────────────────────────────────
    _add_heading(doc, "六、价值投资财务指标风险", level=1)
    risk_items = data.get("financial_risks", [])
    if risk_items:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["风险项", "状态", "说明"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
            _set_cell_bg(hdr[i], "D9E1F2")
        for item in risk_items:
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            status = item.get("status", "")
            row[1].text = status
            _set_cell_bg(row[1], _status_color(status))
            row[2].text = item.get("note", "")
        doc.add_paragraph()

    # ── Section 7: 成长与风险量化 ─────────────────────────────────────────────
    _add_heading(doc, "七、成长与风险量化", level=1)
    growth_risk = data.get("growth_and_risk", "")
    if growth_risk:
        doc.add_paragraph(growth_risk)
    doc.add_paragraph()

    # ── Section 8: 综合结论 ───────────────────────────────────────────────────
    _add_heading(doc, "八、综合结论", level=1)
    conclusion_final = data.get("conclusion", "")
    if conclusion_final:
        p = doc.add_paragraph(conclusion_final)
        p.runs[0].font.bold = True
    doc.add_paragraph()

    # ── Footer disclaimer ─────────────────────────────────────────────────────
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disclaimer.add_run("风险提示：本报告基于公开数据分析，不构成投资建议。投资有风险，入市需谨慎。")
    dr.font.size = Pt(9)
    dr.font.color.rgb = COLOR_GRAY
    dr.font.italic = True

    doc.save(output_path)
    print(f"报告已生成：{output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate DOCX report from MR Dang analysis JSON",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate report from JSON file
  uv run scripts/report.py report.json

  # Specify output path
  uv run scripts/report.py report.json --output 东方电缆分析报告.docx

JSON structure expected:
  {
    "meta": {"stock_name": "东方电缆", "ts_code": "603606.SH", "industry": "电缆",
             "analysis_date": "2026-04-02", "data_source": "Tushare + 年报 + Tavily"},
    "screening": [{"name": "题材股筛查", "result": "通过", "note": "..."}],
    "screening_conclusion": "通过全部筛查",
    "valuation": {"pe_ttm": 18.5, "pb": 2.1, "total_mv": "390亿", "circ_mv": "380亿"},
    "financials": {"dv_ratio": 3.2, "debt_to_assets": 45.0, "roe": 15.3, ...},
    "dividend": {"dividend_count": 3, "avg_cash_div_per_10_shares": 1.2, "dividend_stability": "稳定分红"},
    "business": {"main_business": "...", "industry_position": "...", "cost_advantage": "..."},
    "scores": [{"dimension": "生产资料属性", "score": 18, "full_score": 20, "reason": "..."}],
    "total_score": 78,
    "rating": "⭐⭐⭐⭐ 良好",
    "suggestion": "可分批买入",
    "checklist": [{"name": "三句话逻辑", "status": "达标", "note": "..."}],
    "financial_risks": [{"name": "商誉/资产比率", "status": "达标", "note": "0.5%"}],
    "growth_and_risk": "...",
    "conclusion": "..."
  }
        """,
    )
    parser.add_argument("input", help="Input JSON file path (use - for stdin)")
    parser.add_argument("--output", "-o", help="Output DOCX file path (default: <input_stem>_report.docx)")

    args = parser.parse_args()

    if args.input == "-":
        data = json.load(sys.stdin)
    else:
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"Error: file not found: {input_path}", file=sys.stderr)
            sys.exit(1)
        with open(input_path, encoding="utf-8") as f:
            data = json.load(f)

    if args.output:
        output_path = Path(args.output)
    elif args.input == "-":
        output_path = Path("report.docx")
    else:
        output_path = Path(args.input).with_name(Path(args.input).stem + "_report.docx")

    build_report(data, output_path)


if __name__ == "__main__":
    main()
